from pathlib import Path
import pytest
from docx import Document
from app.core.lrmx import LrmxFile
from app.core.docx_exporter import (
    DocxExporter,
    _format_time6, _format_time8,
    _format_birth, _format_retire_age,
    _format_jianli_list, _calc_age,
)


# ── Unit tests for helper functions ──────────────────────────────────────────

def test_format_time6_converts_yyyymm():
    assert _format_time6('199001') == '1990.01'
    assert _format_time6('202312') == '2023.12'


def test_format_time6_passthrough_empty():
    assert _format_time6('   ') == ''
    assert _format_time6('') == ''


def test_format_time8_converts_yyyymmdd():
    assert _format_time8('20260506') == '2026.05.06'


def test_calc_age_returns_positive():
    assert _calc_age('199001') > 30


def test_format_birth_includes_age():
    result = _format_birth('199001')
    assert '1990.01' in result
    assert '岁' in result
    assert '\n' in result


def test_format_retire_age_calculates_from_birth():
    from app.core.docx_exporter import _format_retire_age
    # born 196911, retires 203102 → 2031-1969=62, month 02<11 → 61 years
    result = _format_retire_age('203102', '196911')
    assert '2031.02' in result
    assert '61岁' in result


def test_format_retire_age_no_birth():
    from app.core.docx_exporter import _format_retire_age
    result = _format_retire_age('203102', '')
    assert '2031.02' in result
    assert '岁' not in result


def test_format_jianli_list_normalizes_6digit():
    text = '199001--199601  某大学学习\n199601--  某单位工作'
    result = _format_jianli_list(text)
    assert result[0] == '1990.01--1996.01\t某大学学习'
    assert result[1] == '1996.01--\t某单位工作'


def test_format_jianli_list_already_formatted():
    text = '1990.01--1996.01  某大学学习\n1996.01--         某单位工作'
    result = _format_jianli_list(text)
    assert result[0] == '1990.01--1996.01\t某大学学习'
    assert result[1] == '1996.01--\t某单位工作'


def test_format_jianli_list_emdash_separator():
    text = '2000.09—2003.06  兰州气象学校学习\n2003.06—2004.06  待分配'
    result = _format_jianli_list(text)
    assert result[0] == '2000.09--2003.06\t兰州气象学校学习'
    assert result[1] == '2003.06--2004.06\t待分配'


def test_format_jianli_list_mixed_separators():
    text = (
        '2014.10--2018.03  应急办公室主任\n'
        '2018.03—2019.11  坡头乡党委副书记\n'
        '2026.02--  民政局副局长'
    )
    result = _format_jianli_list(text)
    assert result[0] == '2014.10--2018.03\t应急办公室主任'
    assert result[1] == '2018.03--2019.11\t坡头乡党委副书记'
    assert result[2] == '2026.02--\t民政局副局长'


def test_format_jianli_list_preserves_internal_dash():
    # The leading em-dash is normalized; the inner '2001.12-2004.12' is kept.
    text = '2000.09—2003.06  学习(期间:2001.12-2004.12自考取得大专学历)'
    result = _format_jianli_list(text)
    assert result[0] == '2000.09--2003.06\t学习(期间:2001.12-2004.12自考取得大专学历)'


def test_format_jianli_list_strips_leading_space():
    text = ' 2021.09--2024.02  政府副县长'
    result = _format_jianli_list(text)
    assert result[0] == '2021.09--2024.02\t政府副县长'


def test_format_jianli_list_passthrough_non_matching():
    text = '这是普通文字\n1990.01--1996.01  某经历'
    result = _format_jianli_list(text)
    assert '\t这是普通文字' in result
    assert '1990.01--1996.01\t某经历' in result


def test_format_jianli_list_returns_list():
    result = _format_jianli_list('201207--201606  经历一\n201607--  经历二')
    assert isinstance(result, list)
    assert len(result) == 2


def test_format_jianli_list_empty():
    assert _format_jianli_list('') == []


# ── Integration tests ─────────────────────────────────────────────────────────

def make_template(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph('姓名：{{XingMing}}')
    doc.add_paragraph('出生年月：{{ChuShengNianYue}}')
    doc.add_paragraph('入党时间：{{RuDangShiJian}}')
    doc.add_paragraph('照片：{{ZhaoPian}}')
    doc.save(path)
    return path


def test_export_fills_basic_fields(sample_lrmx, tmp_path):
    tpl_path = make_template(tmp_path / 'template.docx')
    exporter = DocxExporter(tpl_path)
    out = tmp_path / 'out.docx'
    lf = LrmxFile(sample_lrmx)
    exporter.export(lf, out)
    assert out.exists()
    doc = Document(out)
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    assert '张三' in full_text
    assert '{{' not in full_text


def test_export_formats_time_field(sample_lrmx, tmp_path):
    tpl_path = make_template(tmp_path / 'template.docx')
    exporter = DocxExporter(tpl_path)
    out = tmp_path / 'out.docx'
    lf = LrmxFile(sample_lrmx)
    exporter.export(lf, out)
    doc = Document(out)
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    # RuDangShiJian '201506' → '2015.06'
    assert '2015.06' in full_text


def test_export_formats_age_field(sample_lrmx, tmp_path):
    tpl_path = make_template(tmp_path / 'template.docx')
    exporter = DocxExporter(tpl_path)
    out = tmp_path / 'out.docx'
    lf = LrmxFile(sample_lrmx)
    exporter.export(lf, out)
    doc = Document(out)
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    # ChuShengNianYue '199001' → '1990.01\n（XX岁）'
    assert '1990.01' in full_text
    assert '岁' in full_text


def test_export_photo_empty_string_when_no_photo(sample_lrmx, tmp_path):
    tpl_path = make_template(tmp_path / 'template.docx')
    exporter = DocxExporter(tpl_path)
    out = tmp_path / 'out.docx'
    lf = LrmxFile(sample_lrmx)
    # sample_lrmx has empty ZhaoPian - should not raise
    exporter.export(lf, out)
    assert out.exists()


def test_export_raises_if_template_missing(sample_lrmx, tmp_path):
    exporter = DocxExporter(tmp_path / 'nonexistent.docx')
    with pytest.raises(Exception):
        exporter.export(LrmxFile(sample_lrmx), tmp_path / 'out.docx')


def test_build_context_strips_invisible(sample_lrmx, tmp_path):
    tpl_path = make_template(tmp_path / 'template.docx')
    from docxtpl import DocxTemplate
    tpl = DocxTemplate(tpl_path)
    exporter = DocxExporter(tpl_path)
    lf = LrmxFile(sample_lrmx)
    ctx = exporter._build_context(lf, tpl)
    # XingMing '张三' has no invisible chars, should be '张三'
    assert ctx['XingMing'] == '张三'


def test_build_context_jiating_indexed_slots(sample_lrmx, tmp_path):
    tpl_path = make_template(tmp_path / 'template.docx')
    from docxtpl import DocxTemplate
    tpl = DocxTemplate(tpl_path)
    exporter = DocxExporter(tpl_path)
    ctx = exporter._build_context(LrmxFile(sample_lrmx), tpl)
    # m0 has the one family member from the fixture
    assert ctx['m0']['ChengWei'] == '妻子'
    assert '岁' in ctx['m0']['Age']
    assert ctx['m0']['ChuShengRiQi'] == '1992.05'
    # m1 and beyond are empty dicts
    assert ctx['m1']['ChengWei'] == ''
    assert ctx['m9']['XingMing'] == ''


def test_build_context_jianli_is_list(sample_lrmx, tmp_path):
    tpl_path = make_template(tmp_path / 'template.docx')
    from docxtpl import DocxTemplate
    tpl = DocxTemplate(tpl_path)
    exporter = DocxExporter(tpl_path)
    ctx = exporter._build_context(LrmxFile(sample_lrmx), tpl)
    assert isinstance(ctx['JianLi'], list)
    assert ctx['JianLi'][0] == '2012.07--2016.06\t某大学某专业学习'
    assert ctx['JianLi'][1] == '2016.07--\t某单位科员'


def test_export_jianli_uses_tab(sample_lrmx, tmp_path):
    # Template: a table cell that loops the JianLi list into one paragraph per line.
    # The {%p for %} / {%p endfor %} tags must each occupy their own paragraph.
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.paragraphs[0].text = '{%p for line in JianLi %}'
    cell.add_paragraph('{{line}}')
    cell.add_paragraph('{%p endfor %}')
    tpl_path = tmp_path / 'template.docx'
    doc.save(tpl_path)

    exporter = DocxExporter(tpl_path)
    out = tmp_path / 'out.docx'
    exporter.export(LrmxFile(sample_lrmx), out)

    result = Document(out)
    jianli_paras = [
        p
        for t in result.tables
        for row in t.rows
        for c in row.cells
        for p in c.paragraphs
        if '某大学某专业学习' in p.text or '某单位科员' in p.text
    ]
    assert len(jianli_paras) == 2
    for p in jianli_paras:
        # Time and experience separated by a real Word tab (reported as \t).
        assert '\t' in p.text
        # The code no longer sets indentation; that is configured in the template.
        assert p.paragraph_format.left_indent is None
        assert p.paragraph_format.first_line_indent is None


def test_build_family_deceased_clears_age(tmp_path):
    """Family member with '已去世' or '已故' in GongZuoDanWeiJiZhiWu gets Age=''."""
    lrmx_xml = """<?xml version="1.0" encoding="UTF-8"?>
<Person>
  <XingMing>测试</XingMing>
  <JiaTingChengYuan>
    <Item>
      <ChengWei>父亲</ChengWei>
      <XingMing>张父</XingMing>
      <ChuShengRiQi>194501</ChuShengRiQi>
      <ZhengZhiMianMao>群众</ZhengZhiMianMao>
      <GongZuoDanWeiJiZhiWu>已去世</GongZuoDanWeiJiZhiWu>
    </Item>
    <Item>
      <ChengWei>母亲</ChengWei>
      <XingMing>张母</XingMing>
      <ChuShengRiQi>194803</ChuShengRiQi>
      <ZhengZhiMianMao>群众</ZhengZhiMianMao>
      <GongZuoDanWeiJiZhiWu>已故</GongZuoDanWeiJiZhiWu>
    </Item>
    <Item>
      <ChengWei>妻子</ChengWei>
      <XingMing>张妻</XingMing>
      <ChuShengRiQi>197503</ChuShengRiQi>
      <ZhengZhiMianMao>群众</ZhengZhiMianMao>
      <GongZuoDanWeiJiZhiWu>某单位职工</GongZuoDanWeiJiZhiWu>
    </Item>
  </JiaTingChengYuan>
</Person>"""
    p = tmp_path / 'test.lrmx'
    p.write_text(lrmx_xml, encoding='utf-8')
    tpl_path = make_template(tmp_path / 'template.docx')
    from docxtpl import DocxTemplate
    tpl = DocxTemplate(tpl_path)
    exporter = DocxExporter(tpl_path)
    ctx = exporter._build_context(LrmxFile(p), tpl)
    assert ctx['m0']['Age'] == ''          # 已去世
    assert ctx['m1']['Age'] == ''          # 已故
    assert ctx['m2']['Age'] != ''          # alive → has age


def test_build_context_retire_field(sample_lrmx, tmp_path):
    tpl_path = make_template(tmp_path / 'template.docx')
    from docxtpl import DocxTemplate
    tpl = DocxTemplate(tpl_path)
    exporter = DocxExporter(tpl_path)
    ctx = exporter._build_context(LrmxFile(sample_lrmx), tpl)
    # DaoLingNianYue 205501, ChuShengNianYue 199001 → 65 years old at retirement
    assert '2055.01' in ctx['DaoLingNianYue']
    assert '65岁' in ctx['DaoLingNianYue']
