import base64
import html as _html
import math
import re
import sys
from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Optional

from .lrmx import LrmxFile


def get_template_path() -> Path:
    """返回内置模板路径，开发环境和 PyInstaller 打包后均适用。"""
    if hasattr(sys, '_MEIPASS'):
        p = Path(sys._MEIPASS) / 'resources' / 'template.docx'
    else:
        p = Path(__file__).parent.parent / 'resources' / 'template.docx'
    if not p.exists():
        raise FileNotFoundError(
            f'找不到内置模板文件，请将 template.docx 放至：{p}'
        )
    return p


# ── Constants ─────────────────────────────────────────────────────────────────

MAX_FAMILY_SLOTS = 10   # how many m0..m9 slots to expose; should exceed any template

# JianLi font tiers: (max_line_count, font_size_pt)
# Starting at (12, 14.0), each tier adds 2 lines and drops 0.5 pt.
# Line spacing = font_size + 1 pt.
_JIANLI_FONT_TIERS: list[tuple[int, float]] = [
    (12, 14.0),
    (18, 13.5),
    (24, 13.0),
    (28, 12.5),
    (34, 12.0),
    (38, 11.5),
    (40, 11.0),
    (44, 10.5),
    (46, 10.0),
    (58,  9.5),
    (58,  9.0),
    (58,  8.5),
    (58,  8.0),
    (58,  7.5),
    (58,  7.0),
    (58,  6.5),
    (58,  6.0),
    (58,  5.5),
    (58,  5.0),
]

_PT_PER_EMU = 12700      # 1 pt = 12700 EMU
_CELL_MARGIN_PT = 0    # Word default cell left/right inner margin ≈ 1.9 mm each side
_MIN_FONT_PT = 8.0       # never shrink below 8 pt

_EMPTY_MEMBER: dict[str, str] = {
    'ChengWei': '', 'XingMing': '', 'Age': '',
    'ChuShengRiQi': '', 'ZhengZhiMianMao': '', 'GongZuoDanWeiJiZhiWu': '',
}

# ── Field classification ──────────────────────────────────────────────────────

_TIME6_FIELDS = frozenset({
    'RuDangShiJian', 'CanJiaGongZuoShiJian', 'TianBiaoShiJian',
})
_TIME8_FIELDS = frozenset({'JiSuanNianLingShiJian'})

_XUELI_KEY  = 'QuanRiZhiJiaoYu_XueLi_BiYeYuanXiaoXi'
_XUEWEI_KEY = 'QuanRiZhiJiaoYu_XueWei_BiYeYuanXiaoXi'

# Matches all whitespace variants + invisible Unicode chars
_INVIS = re.compile(r'[\s​‌‍﻿ 　]+')

# JianLi source line: leading start date, any dash/range separator, optional end
# date, then the experience. The separator may be '--', a full-width em/en dash,
# a tilde, '至', etc. Only this LEADING separator is normalized; dashes inside the
# experience text are left untouched.
_JIANLI_LINE = re.compile(
    r'^\s*(\d{4}\.\d{2}|\d{6})\s*(?:[—–－〜～~-]+|至)\s*(\d{4}\.\d{2}|\d{6})?\s*(.*)$'
)
# Pattern that appears in rendered JianLi paragraphs (e.g. "1990.01--1996.06\t...")
_JIANLI_RENDERED = re.compile(r'\d{4}\.\d{2}--')


# ── Pure helper functions ─────────────────────────────────────────────────────

def _to_ym(s: str) -> str:
    s = s.strip() if s else ''
    if re.fullmatch(r'\d{6}', s):
        return f'{s[:4]}.{s[4:]}'
    return s


def _format_time6(value: str) -> str:
    v = _INVIS.sub('', value)
    if re.fullmatch(r'\d{6}', v):
        return f'{v[:4]}.{v[4:]}'
    return v


def _format_time8(value: str) -> str:
    v = _INVIS.sub('', value)
    if re.fullmatch(r'\d{8}', v):
        return f'{v[:4]}.{v[4:6]}.{v[6:]}'
    return v


def _calc_age(ym6: str) -> int:
    """Age in completed years from yyyyMM to today."""
    if not re.fullmatch(r'\d{6}', ym6):
        return 0
    year, month = int(ym6[:4]), int(ym6[4:])
    today = date.today()
    age = today.year - year
    if today.month < month:
        age -= 1
    return age


def _format_birth(value: str) -> str:
    """ChuShengNianYue → 'yyyy.MM\n（X岁）'."""
    v = _INVIS.sub('', value)
    formatted = _format_time6(v)
    if re.fullmatch(r'\d{6}', v):
        return f'{formatted}\n（{_calc_age(v)}岁）'
    return formatted


def _format_retire_age(value: str, birth_ym6: str) -> str:
    """DaoLingNianYue → 'yyyy.MM\n（X岁）' where X is the retirement age."""
    v = _INVIS.sub('', value)
    b = _INVIS.sub('', birth_ym6)
    formatted = _format_time6(v)
    if re.fullmatch(r'\d{6}', v) and re.fullmatch(r'\d{6}', b):
        d_year, d_month = int(v[:4]), int(v[4:])
        b_year, b_month = int(b[:4]), int(b[4:])
        retire_age = d_year - b_year
        if d_month < b_month:
            retire_age -= 1
        return f'{formatted}\n（{retire_age}岁）'
    return formatted


def _format_jianli_list(text: str) -> list[str]:
    """Normalize each JianLi line to 'yyyy.MM--[yyyy.MM]<TAB>experience'.

    Accepts mixed source separators ('--', '—', '–', '~', '至', ...) and
    canonicalizes the leading one to '--'. A tab separates the time part from the
    experience; docxtpl renders it as a real Word tab. Column alignment is left to
    the template: set a hanging indent on the JianLi paragraph so the first line's
    experience snaps to the tab stop and wrapped lines align under it.
    """
    if not text or not text.strip():
        return []
    out = []
    for line in text.split('\n'):
        if not line.strip():
            continue
        m = _JIANLI_LINE.match(line)
        if not m:
            out.append('\t' + line.strip())
            continue
        start = _to_ym(m.group(1))
        end = _to_ym(m.group(2)) if m.group(2) else ''
        content = m.group(3).strip()
        out.append(f'{start}--{end}\t{content}')
    return out


_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _tr_height_emu(tr) -> int:
    """Read <w:trHeight w:val> from a raw <w:tr> element; val is in twips → EMU."""
    tr_pr = tr.find(f'{{{_W}}}trPr')
    if tr_pr is None:
        return 0
    tr_h = tr_pr.find(f'{{{_W}}}trHeight')
    if tr_h is None:
        return 0
    val = tr_h.get(f'{{{_W}}}val')
    return int(val) * 635 if val else 0   # 1 twip = 635 EMU


def _grid_col_of(tr, target_tc) -> int:
    """Return the grid-column index of target_tc inside tr."""
    col = 0
    for tc in tr.findall(f'{{{_W}}}tc'):
        if tc is target_tc:
            return col
        tc_pr = tc.find(f'{{{_W}}}tcPr')
        span = 1
        if tc_pr is not None:
            gs = tc_pr.find(f'{{{_W}}}gridSpan')
            if gs is not None:
                span = int(gs.get(f'{{{_W}}}val', 1))
        col += span
    return col


def _tc_at_grid_col(tr, target_col: int):
    """Return the <w:tc> element at grid-column target_col in tr, or None."""
    col = 0
    for tc in tr.findall(f'{{{_W}}}tc'):
        tc_pr = tc.find(f'{{{_W}}}tcPr')
        span = 1
        if tc_pr is not None:
            gs = tc_pr.find(f'{{{_W}}}gridSpan')
            if gs is not None:
                span = int(gs.get(f'{{{_W}}}val', 1))
        if col <= target_col < col + span:
            return tc
        col += span
    return None


def _shrink_para_by_1pt(para, fallback_pt: float = 10.5, target_pt: float = 13.5) -> bool:
    """Set every run's font size to target_pt. Returns True if anything changed."""
    from docx.shared import Pt
    if not para.runs:
        return False
    font_pt = _para_font_size_pt(para) or fallback_pt
    if target_pt == font_pt:
        return False
    for run in para.runs:
        run.font.size = Pt(target_pt)
    return True


def _char_width_pt(ch: str, font_pt: float) -> float:
    """CJK / full-width chars count as full font_pt wide; ASCII as ~60%."""
    return font_pt if ord(ch) > 0x2E7F else font_pt


def _text_width_pt(text: str, font_pt: float) -> float:
    return sum(_char_width_pt(c, font_pt) for c in text)


def _para_font_size_pt(para) -> Optional[float]:
    """Return font size in pt, walking run → style chain → base styles."""
    for run in para.runs:
        if run.font.size:
            return run.font.size / _PT_PER_EMU
    style = para.style
    while style is not None:
        try:
            if style.font.size:
                return style.font.size / _PT_PER_EMU
        except Exception:
            pass
        style = getattr(style, 'base_style', None)
    return None


# ── Parallel DOCX worker ──────────────────────────────────────────────────────

from concurrent.futures import ProcessPoolExecutor, as_completed as _as_completed
from typing import Callable

_g_docx_template_bytes: bytes = b''
_g_docx_cell_size: 'tuple[int, int]' = (-1, -1)


def _init_docx_worker(template_bytes: bytes, cell_size: 'tuple[int, int]') -> None:
    global _g_docx_template_bytes, _g_docx_cell_size
    _g_docx_template_bytes = template_bytes
    _g_docx_cell_size = cell_size


def _docx_worker(args: 'tuple[str, str]') -> str:
    lrmx_path_str, output_path_str = args
    from pathlib import Path as _Path
    from app.core.lrmx import LrmxFile as _LrmxFile
    lf = _LrmxFile(_Path(lrmx_path_str))
    DocxExporter(_g_docx_template_bytes, _g_docx_cell_size).export(lf, _Path(output_path_str))
    return output_path_str


# ── Exporter ──────────────────────────────────────────────────────────────────

class DocxExporter:
    def __init__(self, template: 'bytes | Path', cell_size: 'tuple[int, int] | None' = None) -> None:
        if isinstance(template, bytes):
            self._template_bytes: bytes = template
        else:
            self._template_bytes = Path(template).read_bytes()
        # cell_size=None → probe lazily on first use
        # cell_size=(w,h) with w>0 → found, use directly
        # cell_size=(-1,-1) → probed externally, not found
        self._photo_cell_cache: tuple[int, int] | None = cell_size
        self._xueli_shrink_text: str | None = None
        self._jianli_line_count: int = 0
        self._jianli_lines: list[str] = []

    def export(self, lrmx: LrmxFile, output_path: Path) -> None:
        from docxtpl import DocxTemplate
        tpl = DocxTemplate(BytesIO(self._template_bytes))
        context = self._build_context(lrmx, tpl)
        tpl.render(context)
        out = Path(output_path)
        tpl.save(out)
        self._post_process(out)

    def _build_context(self, lrmx: LrmxFile, tpl) -> dict:
        raw = lrmx.as_dict()
        birth_ym = _INVIS.sub('', raw.get('ChuShengNianYue', ''))

        ctx: dict = {}
        for key, value in raw.items():
            if key in ('JianLi', _XUELI_KEY, _XUEWEI_KEY):
                continue  # handled separately below
            elif key == 'ZhaoPian':
                ctx[key] = self._decode_photo(value, tpl)
            elif key == 'ChuShengNianYue':
                ctx[key] = _format_birth(value)
            elif key == 'DaoLingNianYue':
                ctx[key] = _format_retire_age(value, birth_ym)
            elif key in _TIME6_FIELDS:
                ctx[key] = _format_time6(value)
            elif key in _TIME8_FIELDS:
                ctx[key] = _format_time8(value)
            else:
                ctx[key] = _html.escape(_INVIS.sub('', value), quote=False)

        # JianLi: list of 'time<TAB>experience' strings for a
        # {%p for line in JianLi %}{{line}}{%p endfor %} loop. Column alignment
        # (hanging indent) is configured in the template, not here.
        jianli_raw = _format_jianli_list(raw.get('JianLi', ''))
        jianli = [_html.escape(l, quote=False) for l in jianli_raw]
        self._jianli_line_count = len(jianli)
        self._jianli_lines = jianli_raw   # unescaped, for width estimation
        ctx['JianLi'] = jianli

        # QuanRiZhiJiaoYu XueLi/XueWei overflow:
        #   > 12 chars + XueWei empty  → split: XueLi[:10] / XueLi[10:] → XueWei
        #   > 12 chars + XueWei filled → keep as-is, flag for post-process font shrink
        self._xueli_shrink_text = None
        xueli  = _INVIS.sub('', raw.get(_XUELI_KEY,  ''))
        xuewei = _INVIS.sub('', raw.get(_XUEWEI_KEY, ''))
        if len(xueli) > 12 and not xuewei:
            ctx[_XUELI_KEY]  = xueli[:10]
            ctx[_XUEWEI_KEY] = xueli[10:]
        elif len(xueli) > 12:
            ctx[_XUELI_KEY]  = xueli
            ctx[_XUEWEI_KEY] = xuewei
            self._xueli_shrink_text = xueli
        else:
            ctx[_XUELI_KEY]  = xueli
            ctx[_XUEWEI_KEY] = xuewei

        # JiaTingChengYuan: fixed indexed slots m0..m(MAX_FAMILY_SLOTS-1)
        family = self._build_family(lrmx)
        for i in range(MAX_FAMILY_SLOTS):
            ctx[f'm{i}'] = family[i] if i < len(family) else dict(_EMPTY_MEMBER)

        return ctx

    # ── Post-processing ───────────────────────────────────────────────────────

    def _post_process(self, path: Path) -> None:
        from docx import Document as DocxDoc
        doc = DocxDoc(path)
        changed = self._shrink_overflow_cells(doc)
        changed |= self._shrink_xueli_cell(doc)
        changed |= self._shrink_jianli_cell(doc)
        if changed:
            doc.save(path)

    def _shrink_overflow_cells(self, doc) -> bool:
        """In rows with exact height, shrink font by 1 pt when text exceeds cell width."""
        from docx.enum.table import WD_ROW_HEIGHT_RULE

        changed = False
        for table in doc.tables:
            seen: set[int] = set()
            for row in table.rows:
                if row.height_rule != WD_ROW_HEIGHT_RULE.EXACTLY:
                    continue
                for cell in row.cells:
                    cid = id(cell._tc)
                    if cid in seen:
                        continue
                    seen.add(cid)
                    if self._shrink_cell_if_needed(cell):
                        changed = True
        return changed

    @staticmethod
    def _shrink_cell_if_needed(cell) -> bool:
        from docx.shared import Pt
        cell_w = cell.width
        if not cell_w:
            return False
        usable_pt = cell_w / _PT_PER_EMU - 2 * _CELL_MARGIN_PT
        if usable_pt <= 0:
            return False
        changed = False
        for para in cell.paragraphs:
            text = para.text
            if not text.strip():
                continue
            font_pt = _para_font_size_pt(para)
            if not font_pt:
                continue
            if _text_width_pt(text, font_pt) > usable_pt:
                new_pt = max(font_pt - 1.0, _MIN_FONT_PT)
                for run in para.runs:
                    run.font.size = Pt(new_pt)
                changed = True
        return changed

    def _shrink_xueli_cell(self, doc) -> bool:
        """Shrink XueLi paragraph by 1 pt when it was too long and XueWei was filled."""
        if not self._xueli_shrink_text:
            return False
        changed = False
        for table in doc.tables:
            seen: set[int] = set()
            for row in table.rows:
                for cell in row.cells:
                    cid = id(cell._tc)
                    if cid in seen:
                        continue
                    seen.add(cid)
                    for para in cell.paragraphs:
                        if para.text == self._xueli_shrink_text:
                            changed |= _shrink_para_by_1pt(para)
        return changed

    # ── 新增：获取简历单元格高度（EMU） ──────────────────────────────────────────

    def _get_cell_height_emu(self, cell) -> int:
        """
        返回 cell 所在行（含 vMerge 合并行）的累计高度（EMU）。
        未设置固定高度时返回 0。
        """
        tc = cell._tc
        tr = tc.getparent()  # <w:tr>
        tbl = tr.getparent()  # <w:tbl>
        trs = tbl.findall(f'{{{_W}}}tr')

        # 确定 cell 在哪一列（grid col）
        grid_col = _grid_col_of(tr, tc)

        # 找到 vMerge restart 所在行（向上找 restart 行）
        start_tr = tr
        tr_idx = trs.index(tr)
        for i in range(tr_idx, -1, -1):
            candidate_tc = _tc_at_grid_col(trs[i], grid_col)
            if candidate_tc is None:
                break
            tc_pr = candidate_tc.find(f'{{{_W}}}tcPr')
            if tc_pr is None:
                start_tr = trs[i]
                break
            vm = tc_pr.find(f'{{{_W}}}vMerge')
            if vm is None:
                # 非合并单元格，就是起点
                start_tr = trs[i]
                break
            if vm.get(f'{{{_W}}}val', '') == 'restart':
                start_tr = trs[i]
                break

        # 从 restart 行开始向下累加高度
        start_idx = trs.index(start_tr)
        total_h = _tr_height_emu(start_tr)
        for next_tr in trs[start_idx + 1:]:
            next_tc = _tc_at_grid_col(next_tr, grid_col)
            if next_tc is None:
                break
            next_pr = next_tc.find(f'{{{_W}}}tcPr')
            next_vm = next_pr.find(f'{{{_W}}}vMerge') if next_pr is not None else None
            if next_vm is not None and next_vm.get(f'{{{_W}}}val', '') != 'restart':
                total_h += _tr_height_emu(next_tr)
            else:
                break
        return total_h

    def _estimate_jianli_visual_lines(self, cell_width_emu: int, font_pt: float) -> int:
        """估算简历单元格视觉行数。

        每条目格式：'yyyy.MM--yyyy.MM\t经历内容'
        悬挂缩进 9 个汉字宽，内容（含首行 tab 后）都从 9*font_pt 处开始，
        可用宽度 = 单元格宽 - 左右内边距 - 悬挂缩进。
        """
        cell_pt = cell_width_emu / _PT_PER_EMU
        margin_pt = _CELL_MARGIN_PT * 2        # 左右内边距合计
        indent_pt = 9 * font_pt                # 悬挂缩进（9个汉字）
        avail_pt = cell_pt - margin_pt - indent_pt

        if avail_pt <= 0:
            return self._jianli_line_count

        total = 0
        for idx, line in enumerate(self._jianli_lines):
            # tab 后面才是正文内容；tab 前的日期部分占据悬挂缩进区域
            content = line.split('\t', 1)[1] if '\t' in line else line
            content_w = _text_width_pt(content, font_pt)
            content_line_count = max(1, math.ceil(content_w / avail_pt))
            total += content_line_count
        return total

    def _shrink_jianli_cell(self, doc) -> bool:
        from docx.shared import Pt
        from docx.enum.text import WD_LINE_SPACING
        from typing import NamedTuple

        class _CellInfo(NamedTuple):
            cell: object
            width_emu: int
            height_pt: float

        # ── 第一步：找到简历单元格，同时取宽度和高度 ──────────────────────────────
        jianli_cells: list[_CellInfo] = []
        seen: set[int] = set()
        for table in doc.tables:
            for row in table.rows:
                row_h_pt = (row.height / _PT_PER_EMU) if row.height else 0.0
                for cell in row.cells:
                    if not any(_JIANLI_RENDERED.search(p.text) for p in cell.paragraphs):
                        continue
                    cid = id(cell._tc)
                    if cid in seen:
                        continue
                    seen.add(cid)
                    jianli_cells.append(_CellInfo(
                        cell=cell,
                        width_emu=cell.width or 0,
                        height_pt=row_h_pt,
                    ))

        if not jianli_cells:
            return False

        cell_w = jianli_cells[0].width_emu
        cell_h_pt = jianli_cells[0].height_pt

        # ── 第二步：枚举字号，求满足 f(x)*(x+1) <= cell_h_pt 的最大 x ────────────
        max_font_pt = _JIANLI_FONT_TIERS[0][1]
        candidates = [
            round(max_font_pt - 0.5 * i, 1)
            for i in range(int((max_font_pt - _MIN_FONT_PT) / 0.5) + 1)
        ]

        target_pt = _MIN_FONT_PT
        target_spacing_pt = _MIN_FONT_PT + 1.0

        if cell_w and cell_h_pt:
            for font_pt in candidates:
                spacing_pt = font_pt + 1.0
                vis = self._estimate_jianli_visual_lines(cell_w, font_pt)
                if vis * spacing_pt <= cell_h_pt:
                    target_pt = font_pt
                    target_spacing_pt = spacing_pt
                    print(
                        f'[JianLi] 数据行={self._jianli_line_count}  字号={font_pt}pt'
                        f'  行距={spacing_pt}pt  估算视觉行={vis}'
                        f'  所需高度={vis * spacing_pt:.1f}pt  单元格高={cell_h_pt:.1f}pt  ✓'
                    )
                    break
                print(
                    f'[JianLi] 字号={font_pt}pt  行距={spacing_pt}pt'
                    f'  估算视觉行={vis}  所需高度={vis * spacing_pt:.1f}pt > {cell_h_pt:.1f}pt  ✗'
                )
        else:
            vis = self._jianli_line_count
            for line_limit, font_pt in _JIANLI_FONT_TIERS:
                if vis <= line_limit:
                    target_pt = font_pt
                    target_spacing_pt = font_pt + 1.0
                    break
            else:
                target_pt = _JIANLI_FONT_TIERS[-1][1]
                target_spacing_pt = target_pt + 1.0

        if target_pt >= max_font_pt:
            return False

        # ── 第三步：应用字号与行距 ────────────────────────────────────────────────
        changed = False
        for info in jianli_cells:
            for para in info.cell.paragraphs:
                if _shrink_para_by_1pt(para, target_pt=target_pt):
                    changed = True
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    para.paragraph_format.line_spacing = Pt(target_spacing_pt)
        return changed    # ── Photo ────────────────────────────────────────────────────────────────

    @staticmethod
    def probe_cell_size(template_bytes: bytes) -> tuple[int, int]:
        """扫描模板，返回照片单元格 (width_emu, height_emu)。
        找不到时返回 (-1, -1)。供批量转换前一次性调用。
        """
        try:
            from docx import Document as DocxDoc
            doc = DocxDoc(BytesIO(template_bytes))
            for table in doc.tables:
                trs = table._tbl.findall(f'{{{_W}}}tr')
                for tr_idx, tr in enumerate(trs):
                    for tc in tr.findall(f'{{{_W}}}tc'):
                        if 'ZhaoPian' not in tc.xml:
                            continue
                        tc_pr = tc.find(f'{{{_W}}}tcPr')
                        vm = tc_pr.find(f'{{{_W}}}vMerge') if tc_pr is not None else None
                        if vm is not None and vm.get(f'{{{_W}}}val', '') != 'restart':
                            continue
                        cell_w = 0
                        if tc_pr is not None:
                            tc_w_el = tc_pr.find(f'{{{_W}}}tcW')
                            if tc_w_el is not None:
                                w_val = tc_w_el.get(f'{{{_W}}}w')
                                w_type = tc_w_el.get(f'{{{_W}}}type', '')
                                if w_val and w_type == 'dxa':
                                    cell_w = int(w_val) * 635
                        if cell_w == 0:
                            for row in table.rows:
                                for cell in row.cells:
                                    if cell._tc is tc and cell.width:
                                        cell_w = int(cell.width)
                                        break
                        grid_col = _grid_col_of(tr, tc)
                        total_h = _tr_height_emu(tr)
                        for next_tr in trs[tr_idx + 1:]:
                            next_tc = _tc_at_grid_col(next_tr, grid_col)
                            if next_tc is None:
                                break
                            next_pr = next_tc.find(f'{{{_W}}}tcPr')
                            next_vm = next_pr.find(f'{{{_W}}}vMerge') if next_pr is not None else None
                            if next_vm is not None and next_vm.get(f'{{{_W}}}val', '') != 'restart':
                                total_h += _tr_height_emu(next_tr)
                            else:
                                break
                        if cell_w > 0 and total_h > 0:
                            return (cell_w, total_h)
        except Exception:
            pass
        return (-1, -1)

    @staticmethod
    def export_parallel(
        jobs: 'list[tuple[str, str]]',
        template_bytes: bytes,
        cell_size: 'tuple[int, int]',
        on_progress: 'Callable[[str, str | None, str], None] | None' = None,
    ) -> None:
        """Parallel DOCX export. jobs = [(lrmx_path_str, output_path_str), ...]"""
        if not jobs:
            return
        import os
        n = min(os.cpu_count() or 1, len(jobs))
        with ProcessPoolExecutor(
            max_workers=n,
            initializer=_init_docx_worker,
            initargs=(template_bytes, cell_size),
        ) as executor:
            future_to_stem = {
                executor.submit(_docx_worker, job): Path(job[1]).stem
                for job in jobs
            }
            for future in _as_completed(future_to_stem):
                stem = future_to_stem[future]
                try:
                    output_path = future.result()
                    if on_progress:
                        on_progress(stem, output_path, '')
                except Exception as e:
                    if on_progress:
                        on_progress(stem, None, str(e))

    def _get_photo_cell_size(self) -> Optional[tuple[int, int]]:
        """返回照片单元格 (width_emu, height_emu)，找不到返回 None。"""
        if self._photo_cell_cache is not None:
            return self._photo_cell_cache if self._photo_cell_cache[0] > 0 else None
        result = DocxExporter.probe_cell_size(self._template_bytes)
        self._photo_cell_cache = result
        return result if result[0] > 0 else None

    def _decode_photo(self, b64: str, tpl) -> 'InlineImage | str':
        from docxtpl import InlineImage
        from docx.shared import Emu, Mm
        b64 = _INVIS.sub('', b64)
        if not b64:
            return ''
        # Strip data URI prefix, e.g. "data:image/png;base64,"
        if ',' in b64:
            b64 = b64.split(',', 1)[1]
        try:
            from PIL import Image
            img_bytes = base64.b64decode(b64)
            size = self._get_photo_cell_size()
            if size:
                cell_w_emu, cell_h_emu = size
                # Crop left/right to match cell aspect ratio, then fit by height
                img = Image.open(BytesIO(img_bytes))
                if img.mode not in ('RGB', 'L'):
                    img = img.convert('RGB')
                pw, ph = img.size
                target_w = int(ph * cell_w_emu / cell_h_emu)
                if target_w < pw:
                    left = (pw - target_w) // 2
                    img = img.crop((left, 0, left + target_w, ph))
                buf = BytesIO()
                img.save(buf, format='JPEG')
                buf.seek(0)
                return InlineImage(tpl, buf, height=Emu(cell_h_emu))
            return InlineImage(tpl, BytesIO(img_bytes), height=Mm(45))
        except Exception:
            return ''

    # ── Family members ───────────────────────────────────────────────────────

    def _build_family(self, lrmx: LrmxFile) -> list[dict]:
        members = []
        for raw in lrmx.family_members():
            member: dict = {}
            for k, v in raw.items():
                if k == 'ChuShengRiQi':
                    v_clean = _INVIS.sub('', v)
                    member['Age'] = (
                        f'{_calc_age(v_clean)}岁'
                        if re.fullmatch(r'\d{6}', v_clean) else ''
                    )
                    member[k] = _format_time6(v_clean)
                else:
                    member[k] = _INVIS.sub('', v)
            gongzuo = member.get('GongZuoDanWeiJiZhiWu', '')
            if '已去世' in gongzuo or '已故' in gongzuo:
                member['Age'] = ''
            members.append(member)
        return members
